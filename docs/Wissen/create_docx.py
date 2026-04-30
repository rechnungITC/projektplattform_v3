import sys
import subprocess

# Install python-docx if not available
subprocess.check_call(["sudo", sys.executable, "-m", "pip", "install", "python-docx"])

from docx import Document

doc = Document()
doc.add_heading('Tech-Design: OpenProject Reference Patterns', 0)

with open('/home/ubuntu/openproject_tech_design.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line:
        continue
    if line.startswith('# '):
        pass
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('* '):
        doc.add_paragraph(line[2:], style='List Bullet')
    else:
        doc.add_paragraph(line)

doc.save('/home/ubuntu/OpenProject_Reference_Patterns.docx')
print("Docx created successfully.")
