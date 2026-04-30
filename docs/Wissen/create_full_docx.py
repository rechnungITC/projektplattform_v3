from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ===== STYLES =====
styles = doc.styles

# Title
doc.add_heading('Tech-Design: OpenProject Reference Patterns', 0)

# Subtitle / Meta
p = doc.add_paragraph()
p.add_run('Quelle: ').bold = True
p.add_run('https://github.com/opf/openproject (Branch: dev)')
p.add_run('\n')
p.add_run('Analysiert am: ').bold = True
p.add_run('01. Mai 2026')
p.add_run('\n')
p.add_run('Zweck: ').bold = True
p.add_run('OpenProject Reference Patterns als Grundlage für das Tech-Design')

doc.add_paragraph()

# ===== INHALTSVERZEICHNIS =====
doc.add_heading('Inhaltsverzeichnis', level=1)
toc_items = [
    '1. Repository-Struktur',
    '2. Work Packages (Arbeitspakete)',
    '3. Relations (Beziehungen)',
    '4. Project Hierarchy (Projekthierarchie)',
    '5. Types per Project (Typen pro Projekt)',
    '6. Custom Fields & Attribute-Constraints',
    '7. Scheduling & Derived Dates',
    '8. OpenProject Reference Patterns – Tech-Design-Sektion',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===== 1. REPOSITORY-STRUKTUR =====
doc.add_heading('1. Repository-Struktur', level=1)
doc.add_paragraph(
    'Das OpenProject-Repository (https://github.com/opf/openproject) folgt der Standard-Rails-Konvention. '
    'Die für das Tech-Design relevanten Pfade sind:'
)

# Tabelle: Verzeichnisstruktur
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Pfad'
hdr[1].text = 'Inhalt'
hdr[2].text = 'Relevanz'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

rows = [
    ('app/models/work_package.rb', 'Haupt-Model für Arbeitspakete', 'Hoch'),
    ('app/models/work_package/', 'Concerns: Ancestors, Validations, SchedulingRules, StatusTransitions', 'Hoch'),
    ('app/models/work_packages/', 'Concerns: Relations, DerivedDates, Scopes', 'Hoch'),
    ('app/models/relation.rb', 'Modell für WP-Beziehungen (Relation-Typen)', 'Hoch'),
    ('app/models/project.rb', 'Haupt-Model für Projekte', 'Hoch'),
    ('app/models/projects/', 'Concerns: Hierarchy, Types, AncestorsFromRoot, Versions', 'Hoch'),
    ('app/models/type.rb', 'Modell für Work Package Typen', 'Hoch'),
    ('app/models/type/', 'Concerns: Attributes, AttributeGroups', 'Mittel'),
    ('app/models/work_package_types/', 'Pattern-Resolver für Typ-Konfiguration', 'Mittel'),
    ('db/migrate/', 'Datenbankmigrationen', 'Referenz'),
]
for row_data in rows:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()

# ===== 2. WORK PACKAGES =====
doc.add_heading('2. Work Packages (Arbeitspakete)', level=1)
doc.add_paragraph(
    'Work Packages sind das zentrale Konzept in OpenProject. Sie repräsentieren Aufgaben, Meilensteine, '
    'Phasen, Fehler und alle weiteren konfigurierbaren Arbeitseinheiten. Das Modell WorkPackage '
    '(app/models/work_package.rb) ist durch zahlreiche Concerns modularisiert.'
)

doc.add_heading('2.1 Kern-Assoziationen', level=2)
doc.add_paragraph(
    'Jedes Work Package ist zwingend mit folgenden Entitäten verknüpft:'
)
assoc_table = doc.add_table(rows=1, cols=3)
assoc_table.style = 'Table Grid'
hdr2 = assoc_table.rows[0].cells
hdr2[0].text = 'Assoziation'
hdr2[1].text = 'Typ'
hdr2[2].text = 'Pflichtfeld'
for cell in hdr2:
    cell.paragraphs[0].runs[0].bold = True

assoc_rows = [
    ('belongs_to :project', 'Project', 'Ja'),
    ('belongs_to :type', 'Type', 'Ja'),
    ('belongs_to :status', 'Status', 'Ja'),
    ('belongs_to :author', 'User', 'Ja'),
    ('belongs_to :priority', 'IssuePriority', 'Ja'),
    ('belongs_to :assigned_to', 'Principal', 'Nein'),
    ('belongs_to :responsible', 'Principal', 'Nein'),
    ('belongs_to :version', 'Version', 'Nein'),
    ('belongs_to :category', 'Category', 'Nein'),
    ('has_many :relations', 'Relation (bidirektional)', 'Nein'),
    ('has_many :time_entries', 'TimeEntry', 'Nein'),
    ('has_closure_tree', 'WorkPackageHierarchy (Closure Table)', 'Intern'),
]
for row_data in assoc_rows:
    row = assoc_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()

doc.add_heading('2.2 Modularisierung durch Concerns', level=2)
doc.add_paragraph(
    'Das WorkPackage-Modell inkludiert folgende Concerns:'
)
concerns_table = doc.add_table(rows=1, cols=3)
concerns_table.style = 'Table Grid'
hdr3 = concerns_table.rows[0].cells
hdr3[0].text = 'Concern'
hdr3[1].text = 'Datei'
hdr3[2].text = 'Funktion'
for cell in hdr3:
    cell.paragraphs[0].runs[0].bold = True

concerns_rows = [
    ('WorkPackage::Ancestors', 'app/models/work_package/ancestors.rb', 'Sichtbare Vorfahren im Hierarchiebaum abfragen (Closure Table)'),
    ('WorkPackage::Validations', 'app/models/work_package/validations.rb', 'Pflichtfeld-Validierungen (project, type, author, status, priority)'),
    ('WorkPackage::SchedulingRules', 'app/models/work_package/scheduling_rules.rb', 'Automatische vs. manuelle Terminplanung, soonest_start'),
    ('WorkPackage::StatusTransitions', 'app/models/work_package/status_transitions.rb', 'Statusübergänge und Workflow-Logik'),
    ('WorkPackages::Relations', 'app/models/work_packages/relations.rb', 'has_many-Assoziationen für alle Relationstypen'),
    ('WorkPackages::DerivedDates', 'app/models/work_packages/derived_dates.rb', 'Abgeleitete Termine aus Kindpaketen'),
    ('WorkPackage::CustomActioned', 'app/models/work_package/custom_actioned.rb', 'Custom Actions auf Work Packages'),
]
for row_data in concerns_rows:
    row = concerns_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()

doc.add_heading('2.3 Hierarchie: Closure Table Pattern', level=2)
doc.add_paragraph(
    'Die Parent-Child-Beziehung zwischen Work Packages wird über das Closure Table Pattern implementiert '
    '(has_closure_tree). Die Tabelle WorkPackageHierarchy speichert alle Vorfahren-Nachkommen-Paare mit '
    'der jeweiligen Anzahl von Generationen. Dies ermöglicht effiziente Abfragen über beliebig tiefe '
    'Hierarchien ohne rekursive SQL-Abfragen.'
)
doc.add_paragraph(
    'Das Concern WorkPackage::Ancestors nutzt diese Tabelle, um sichtbare Vorfahren eines Work Packages '
    'unter Berücksichtigung der Benutzerberechtigungen zu ermitteln:'
)
p = doc.add_paragraph()
p.add_run('WorkPackageHierarchy\n').bold = True
p.add_run('  .where(descendant_id: @ids)\n')
p.add_run('  .includes(:ancestor)\n')
p.add_run('  .where(ancestor: { project_id: Project.allowed_to(user, :view_work_packages) })\n')
p.add_run('  .where("generations > 0")\n')
p.add_run('  .order(generations: :desc)\n')
p.add_run('  .group_by(&:descendant_id)')
p.style = doc.styles['No Spacing']

doc.add_paragraph()

# ===== 3. RELATIONS =====
doc.add_heading('3. Relations (Beziehungen)', level=1)
doc.add_paragraph(
    'Das Modell Relation (app/models/relation.rb) bildet alle expliziten Beziehungen zwischen Work Packages ab. '
    'Die Relation hat immer eine Richtung: from → to. Für symmetrische Typen (z. B. relates) wird der Typ '
    'beim Speichern automatisch in die kanonische Richtung umgekehrt (reverse_if_needed).'
)

doc.add_heading('3.1 Relationstypen', level=2)
rel_table = doc.add_table(rows=1, cols=4)
rel_table.style = 'Table Grid'
hdr4 = rel_table.rows[0].cells
hdr4[0].text = 'Konstante'
hdr4[1].text = 'String-Wert'
hdr4[2].text = 'Symmetrisch'
hdr4[3].text = 'Reihenfolge (UI)'
for cell in hdr4:
    cell.paragraphs[0].runs[0].bold = True

rel_rows = [
    ('TYPE_RELATES', '"relates"', 'Ja', '1'),
    ('TYPE_FOLLOWS', '"follows"', 'Nein (Sym: precedes)', '7'),
    ('TYPE_PRECEDES', '"precedes"', 'Nein (Sym: follows)', '6'),
    ('TYPE_BLOCKS', '"blocks"', 'Nein (Sym: blocked)', '4'),
    ('TYPE_BLOCKED', '"blocked"', 'Nein (Sym: blocks)', '5'),
    ('TYPE_DUPLICATES', '"duplicates"', 'Nein (Sym: duplicated)', '6'),
    ('TYPE_DUPLICATED', '"duplicated"', 'Nein (Sym: duplicates)', '7'),
    ('TYPE_INCLUDES', '"includes"', 'Nein (Sym: partof)', '8'),
    ('TYPE_PARTOF', '"partof"', 'Nein (Sym: includes)', '9'),
    ('TYPE_REQUIRES', '"requires"', 'Nein (Sym: required)', '10'),
    ('TYPE_REQUIRED', '"required"', 'Nein (Sym: requires)', '11'),
    ('TYPE_PARENT', '"parent"', 'Virtuell (Closure Table)', '-'),
    ('TYPE_CHILD', '"child"', 'Virtuell (Closure Table)', '-'),
]
for row_data in rel_rows:
    row = rel_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()

doc.add_heading('3.2 Validierungsregeln für Relationen', level=2)
doc.add_paragraph(
    'Der Scope WorkPackages::Scopes::Relatable (app/models/work_packages/scopes/relatable.rb) implementiert '
    'die komplexen Regeln, welche Work Packages als Ziel einer neuen Relation in Frage kommen:'
)
rules = [
    'Kein Zirkel: Relationen dürfen keine Kreise bilden (weder direkt noch transitiv).',
    'Einzelne Relation: Zwischen zwei Work Packages kann nur eine Relation existieren, unabhängig vom Typ.',
    'Kein Vorfahre/Nachkomme: Relationen zwischen direkten oder transitiven Vorfahren/Nachkommen sind nicht erlaubt.',
    'Kein Baum-übergreifender Zirkel: Die Vorfahren-Nachkommen-Kette wird bidirektional betrachtet.',
    'Ausnahme TYPE_RELATES: Nur die "Einzelne Relation"-Regel gilt.',
    'Ausnahme TYPE_PARENT/CHILD: Aktuelle Vorfahren (außer direktem Elternteil) sind erlaubt; Nachkommen nicht.',
]
for rule in rules:
    doc.add_paragraph(rule, style='List Bullet')

doc.add_paragraph()

doc.add_heading('3.3 Scheduling-Relation (follows/precedes)', level=2)
doc.add_paragraph(
    'Die follows-Relation ist die einzige Relation, die direkte Auswirkungen auf die automatische Terminplanung hat. '
    'Das Feld lag (in Arbeitstagen) ermöglicht einen Puffer zwischen Vorgänger und Nachfolger. '
    'Die Methode successor_soonest_start berechnet den frühestmöglichen Starttermin des Nachfolgers:'
)
p = doc.add_paragraph()
p.add_run('def successor_soonest_start\n').bold = True
p.add_run('  if follows? && predecessor_date\n')
p.add_run('    days = WorkPackages::Shared::WorkingDays.new\n')
p.add_run('    days.with_lag(predecessor_date, lag)\n')
p.add_run('  end\n')
p.add_run('end')
p.style = doc.styles['No Spacing']

doc.add_paragraph()

# ===== 4. PROJECT HIERARCHY =====
doc.add_heading('4. Project Hierarchy (Projekthierarchie)', level=1)
doc.add_paragraph(
    'Projekte in OpenProject können hierarchisch strukturiert werden. Das Concern Projects::Hierarchy '
    '(app/models/projects/hierarchy.rb) implementiert das Nested Set Pattern (acts_as_nested_set) '
    'mit den Spalten lft und rgt für effiziente Baumabfragen.'
)

doc.add_heading('4.1 Workspace-Typen', level=2)
doc.add_paragraph(
    'Seit einer neueren Version unterstützt OpenProject drei Workspace-Typen, die die Hierarchiestruktur '
    'auf Projektebene definieren:'
)
ws_table = doc.add_table(rows=1, cols=3)
ws_table.style = 'Table Grid'
hdr5 = ws_table.rows[0].cells
hdr5[0].text = 'Workspace-Typ'
hdr5[1].text = 'Erlaubte Eltern-Typen'
hdr5[2].text = 'Beschreibung'
for cell in hdr5:
    cell.paragraphs[0].runs[0].bold = True

ws_rows = [
    ('project', 'portfolio, program, project', 'Standard-Projekt; kann unter allen anderen liegen'),
    ('program', 'portfolio', 'Programm; kann nur unter einem Portfolio liegen'),
    ('portfolio', '(keine)', 'Portfolio; ist immer Root-Level'),
]
for row_data in ws_rows:
    row = ws_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()

doc.add_heading('4.2 Nested Set Implementierung', level=2)
doc.add_paragraph(
    'Die Projekthierarchie nutzt das Nested Set Pattern. Jedes Projekt hat die Spalten lft und rgt, '
    'die den Bereich im Baum definieren. Alle Nachkommen eines Projekts haben lft-Werte zwischen '
    'dem lft und rgt des Elternprojekts.'
)
doc.add_paragraph(
    'Wichtige Methoden aus Projects::Hierarchy:'
)
methods_table = doc.add_table(rows=1, cols=2)
methods_table.style = 'Table Grid'
hdr6 = methods_table.rows[0].cells
hdr6[0].text = 'Methode'
hdr6[1].text = 'Funktion'
for cell in hdr6:
    cell.paragraphs[0].runs[0].bold = True

methods_rows = [
    ('build_projects_hierarchy(projects)', 'Baut eine verschachtelte Hash-Struktur aus einer flachen Projektliste auf'),
    ('project_tree(projects)', 'Iteriert über den Projektbaum mit Level-Information'),
    ('hierarchy', 'Gibt alle Projekte in der Hierarchie (Eltern + Kinder) zurück'),
    ('has_subprojects?', 'Prüft ob Unterprojekte vorhanden sind'),
    ('active_subprojects', 'Gibt alle aktiven Unterprojekte zurück'),
    ('with_subprojects(bool)', 'Erstellt SQL-Bedingung für Abfragen mit/ohne Unterprojekte'),
    ('ancestors_from_root', 'Gibt Vorfahren von Root bis zum aktuellen Projekt zurück (aufsteigend)'),
]
for row_data in methods_rows:
    row = methods_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()

# ===== 5. TYPES PER PROJECT =====
doc.add_heading('5. Types per Project (Typen pro Projekt)', level=1)
doc.add_paragraph(
    'Die Konfiguration, welche Work Package Typen in einem bestimmten Projekt verfügbar sind, '
    'wird über eine Many-to-Many-Beziehung zwischen Type und Project gelöst.'
)

doc.add_heading('5.1 Datenmodell', level=2)
doc.add_paragraph(
    'Das Modell Type (app/models/type.rb) hat folgende Kern-Assoziationen:'
)
type_table = doc.add_table(rows=1, cols=3)
type_table.style = 'Table Grid'
hdr7 = type_table.rows[0].cells
hdr7[0].text = 'Assoziation'
hdr7[1].text = 'Typ'
hdr7[2].text = 'Beschreibung'
for cell in hdr7:
    cell.paragraphs[0].runs[0].bold = True

type_rows = [
    ('has_and_belongs_to_many :projects', 'HABTM', 'Verbindet Typen mit Projekten (Join-Tabelle: projects_types)'),
    ('has_many :work_packages', 'has_many', 'Alle WPs dieses Typs'),
    ('has_many :workflows', 'has_many', 'Statusübergänge für diesen Typ'),
    ('has_and_belongs_to_many :custom_fields', 'HABTM', 'Custom Fields für diesen Typ (Join-Tabelle: custom_fields_types)'),
    ('belongs_to :color', 'belongs_to', 'Farbe für die UI-Darstellung'),
]
for row_data in type_rows:
    row = type_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()

doc.add_heading('5.2 Aktivierungslogik', level=2)
doc.add_paragraph(
    'Die Aktivierung eines Typs in einem Projekt erfolgt über die HABTM-Beziehung. '
    'Folgende Methoden steuern die Abfrage:'
)
act_table = doc.add_table(rows=1, cols=3)
act_table.style = 'Table Grid'
hdr8 = act_table.rows[0].cells
hdr8[0].text = 'Methode'
hdr8[1].text = 'Klasse'
hdr8[2].text = 'Funktion'
for cell in hdr8:
    cell.paragraphs[0].runs[0].bold = True

act_rows = [
    ('Type.enabled_in(project)', 'Type (class)', 'Gibt alle Typen zurück, die im Projekt aktiviert sind'),
    ('type.enabled_in?(project)', 'Type (instance)', 'Prüft ob dieser Typ im Projekt aktiviert ist'),
    ('project.types', 'Project', 'Gibt alle aktivierten Typen des Projekts zurück (ordered by position)'),
    ('project.rolled_up_types', 'Projects::Types', 'Gibt alle Typen des Projekts und seiner aktiven Unterprojekte zurück'),
    ('project.types_used_by_work_packages', 'Projects::Types', 'Gibt nur die Typen zurück, die tatsächlich von WPs verwendet werden'),
]
for row_data in act_rows:
    row = act_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()

# ===== 6. CUSTOM FIELDS & ATTRIBUTE CONSTRAINTS =====
doc.add_heading('6. Custom Fields & Attribute-Constraints', level=1)
doc.add_paragraph(
    'Das Concern Type::Attributes (app/models/type/attributes.rb) steuert dynamisch, welche Felder '
    '(inklusive Custom Fields) für einen Typ in einem bestimmten Projekt sichtbar sind.'
)

doc.add_heading('6.1 Attribut-Verfügbarkeit', level=2)
doc.add_paragraph(
    'Die Methode passes_attribute_constraint? prüft für jedes Attribut, ob es im aktuellen Kontext '
    '(Typ + Projekt) verfügbar ist:'
)
doc.add_paragraph(
    '1. Für Custom Fields: Ist das Custom Field im Projekt aktiviert? '
    '(custom_field_in_project? prüft project.all_work_package_custom_fields)',
    style='List Number'
)
doc.add_paragraph(
    '2. Für andere Attribute: Gibt es eine registrierte Constraint-Funktion? '
    '(Plugins können Constraints registrieren)',
    style='List Number'
)

doc.add_heading('6.2 Custom Field Caching', level=2)
doc.add_paragraph(
    'Zur Performance-Optimierung werden Custom Fields per Request gecacht. '
    'Der Cache-Key ist project_id + type_id. Die Methode '
    'WorkPackage.available_custom_fields nutzt RequestStore für einen '
    'request-scoped Cache.'
)

doc.add_paragraph()

# ===== 7. SCHEDULING =====
doc.add_heading('7. Scheduling & Derived Dates', level=1)
doc.add_paragraph(
    'OpenProject unterstützt zwei Scheduling-Modi für Work Packages:'
)
sched_table = doc.add_table(rows=1, cols=3)
sched_table.style = 'Table Grid'
hdr9 = sched_table.rows[0].cells
hdr9[0].text = 'Modus'
hdr9[1].text = 'Methode'
hdr9[2].text = 'Verhalten'
for cell in hdr9:
    cell.paragraphs[0].runs[0].bold = True

sched_rows = [
    ('Automatisch', 'schedule_automatically?', 'Termine werden aus Relationen und Kindpaketen abgeleitet'),
    ('Manuell', 'schedule_manually?', 'Termine werden manuell gesetzt und nicht automatisch angepasst'),
]
for row_data in sched_rows:
    row = sched_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
doc.add_paragraph(
    'Das Concern WorkPackages::DerivedDates (app/models/work_packages/derived_dates.rb) berechnet '
    'abgeleitete Termine (derived_start_date, derived_due_date) aus den Kindpaketen. '
    'Diese werden im Scope WorkPackages::Scopes::IncludeDerivedDates eingebunden.'
)

doc.add_paragraph()

# ===== 8. REFERENCE PATTERNS =====
doc.add_heading('8. OpenProject Reference Patterns – Tech-Design-Sektion', level=1)
doc.add_paragraph(
    'Dieser Abschnitt fasst die aus der Code-Analyse gewonnenen Erkenntnisse als '
    '"OpenProject Reference Patterns" zusammen, die direkt ins Tech-Design einfließen.'
)

doc.add_heading('8.1 Pattern: Work Package als Kern-Entität', level=2)
doc.add_paragraph(
    'Work Packages sind die zentrale Arbeitseinheit in OpenProject. Jedes Work Package ist '
    'immer an genau ein Projekt und genau einen Typ gebunden. Diese Bindung ist nicht optional '
    '(validates :project, presence: true; validates :type, presence: true). '
    'Daraus folgt für das Tech-Design:'
)
doc.add_paragraph(
    'Alle Arbeitspakete müssen einem Projekt zugeordnet sein. Eine projektübergreifende '
    'Darstellung ist nur über Queries mit entsprechenden Berechtigungen möglich.',
    style='List Bullet'
)
doc.add_paragraph(
    'Der Typ eines Work Packages muss im zugehörigen Projekt aktiviert sein.',
    style='List Bullet'
)
doc.add_paragraph(
    'Custom Fields sind nur sichtbar, wenn sie sowohl am Typ als auch im Projekt aktiviert sind.',
    style='List Bullet'
)

doc.add_heading('8.2 Pattern: Hierarchie durch Closure Table', level=2)
doc.add_paragraph(
    'Die Work Package Hierarchie (Parent-Child) wird über ein Closure Table Pattern '
    '(WorkPackageHierarchy) implementiert. Dieses Muster ermöglicht:'
)
doc.add_paragraph(
    'Effiziente Abfragen aller Vorfahren oder Nachkommen ohne rekursive SQL-Abfragen.',
    style='List Bullet'
)
doc.add_paragraph(
    'Berechtigungsgefiltertes Traversieren der Hierarchie (visible_ancestors).',
    style='List Bullet'
)
doc.add_paragraph(
    'Abgeleitete Termine (derived_start_date, derived_due_date) aus Kindpaketen.',
    style='List Bullet'
)
doc.add_paragraph(
    'Wichtig: Die Parent-Child-Relation wird NICHT in der Relation-Tabelle gespeichert, '
    'sondern separat in WorkPackageHierarchy. Im Code wird sie jedoch als TYPE_PARENT / '
    'TYPE_CHILD abstrahiert, um eine einheitliche API zu bieten.',
    style='List Bullet'
)

doc.add_heading('8.3 Pattern: Relationstypen und ihre Semantik', level=2)
doc.add_paragraph(
    'OpenProject unterscheidet 11 explizite Relationstypen plus die virtuellen Typen '
    'parent/child. Für das Tech-Design relevant sind insbesondere:'
)
pattern_rel_table = doc.add_table(rows=1, cols=3)
pattern_rel_table.style = 'Table Grid'
hdr10 = pattern_rel_table.rows[0].cells
hdr10[0].text = 'Relationstyp'
hdr10[1].text = 'Semantik'
hdr10[2].text = 'Tech-Design-Implikation'
for cell in hdr10:
    cell.paragraphs[0].runs[0].bold = True

pattern_rel_rows = [
    ('follows / precedes', 'Zeitliche Abfolge mit optionalem Lag (Arbeitstage)', 'Beeinflusst automatische Terminplanung; Lag-Feld beachten'),
    ('blocks / blocked', 'Blockierende Abhängigkeit', 'Blockierte WPs können nicht abgeschlossen werden'),
    ('includes / partof', 'Inhaltliche Zugehörigkeit (nicht hierarchisch)', 'Semantisch ähnlich parent/child, aber kein Scheduling-Effekt'),
    ('duplicates / duplicated', 'Duplikat-Beziehung', 'Beim Schließen eines WP werden Duplikate automatisch geschlossen'),
    ('relates', 'Allgemeine Beziehung', 'Keine spezifische Semantik; nur Einzelrelation-Regel gilt'),
    ('parent / child', 'Hierarchische Eltern-Kind-Beziehung', 'Closure Table; beeinflusst derived_dates und Fortschritt'),
]
for row_data in pattern_rel_rows:
    row = pattern_rel_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()

doc.add_heading('8.4 Pattern: Projekthierarchie mit Workspace-Typen', level=2)
doc.add_paragraph(
    'OpenProject unterstützt eine dreistufige Projekthierarchie: Portfolio → Program → Project. '
    'Diese Hierarchie ist durch ALLOWED_PARENT_WORKSPACE_TYPES strikt geregelt. '
    'Für das Tech-Design bedeutet dies:'
)
doc.add_paragraph(
    'Portfolios sind immer Root-Level und können keine Eltern haben.',
    style='List Bullet'
)
doc.add_paragraph(
    'Programme können nur unter Portfolios liegen.',
    style='List Bullet'
)
doc.add_paragraph(
    'Projekte können unter Portfolios, Programmen oder anderen Projekten liegen.',
    style='List Bullet'
)
doc.add_paragraph(
    'Die Methode rolled_up_types ermöglicht die Aggregation aller verfügbaren Typen '
    'über die gesamte Projekthierarchie.',
    style='List Bullet'
)

doc.add_heading('8.5 Pattern: Types-per-Project Matrix', level=2)
doc.add_paragraph(
    'Die Konfiguration von Typen pro Projekt folgt einer Matrix-Logik:'
)
matrix_table = doc.add_table(rows=1, cols=3)
matrix_table.style = 'Table Grid'
hdr11 = matrix_table.rows[0].cells
hdr11[0].text = 'Ebene'
hdr11[1].text = 'Konfiguration'
hdr11[2].text = 'Abfragemethod'
for cell in hdr11:
    cell.paragraphs[0].runs[0].bold = True

matrix_rows = [
    ('Global', 'Typen werden global definiert (Type.all)', 'Type.all, Type.without_standard'),
    ('Projekt', 'Typen werden pro Projekt aktiviert (HABTM)', 'project.types, Type.enabled_in(project)'),
    ('Typ + Projekt', 'Custom Fields sind pro Typ UND pro Projekt aktivierbar', 'type.passes_attribute_constraint?(attr, project:)'),
    ('Instanz', 'Jedes WP hat genau einen Typ aus den Projekt-Typen', 'work_package.type'),
]
for row_data in matrix_rows:
    row = matrix_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()

doc.add_heading('8.6 Pattern: Berechtigungsgefiltertes Traversieren', level=2)
doc.add_paragraph(
    'Alle Hierarchie-Traversierungen in OpenProject sind berechtigungsgefiltert. '
    'Dies gilt sowohl für die Projekthierarchie als auch für die Work Package Hierarchie:'
)
doc.add_paragraph(
    'visible_ancestors(user): Gibt nur Vorfahren zurück, auf die der Benutzer Zugriff hat.',
    style='List Bullet'
)
doc.add_paragraph(
    'Project.allowed_to(user, :view_work_packages): Filtert Projekte nach Berechtigung.',
    style='List Bullet'
)
doc.add_paragraph(
    'WorkPackage.visible(user): Scope für sichtbare Work Packages.',
    style='List Bullet'
)
doc.add_paragraph(
    'relations.visible(user): Filtert Relationen nach Sichtbarkeit der beteiligten WPs.',
    style='List Bullet'
)

doc.add_paragraph()

# ===== REFERENZEN =====
doc.add_heading('Referenzen (Quelldateien)', level=1)
ref_table = doc.add_table(rows=1, cols=2)
ref_table.style = 'Table Grid'
hdr12 = ref_table.rows[0].cells
hdr12[0].text = 'Datei'
hdr12[1].text = 'GitHub-URL'
for cell in hdr12:
    cell.paragraphs[0].runs[0].bold = True

ref_rows = [
    ('app/models/relation.rb', 'https://github.com/opf/openproject/blob/dev/app/models/relation.rb'),
    ('app/models/work_package.rb', 'https://github.com/opf/openproject/blob/dev/app/models/work_package.rb'),
    ('app/models/work_package/ancestors.rb', 'https://github.com/opf/openproject/blob/dev/app/models/work_package/ancestors.rb'),
    ('app/models/work_package/scheduling_rules.rb', 'https://github.com/opf/openproject/blob/dev/app/models/work_package/scheduling_rules.rb'),
    ('app/models/work_packages/relations.rb', 'https://github.com/opf/openproject/blob/dev/app/models/work_packages/relations.rb'),
    ('app/models/work_packages/scopes/relatable.rb', 'https://github.com/opf/openproject/blob/dev/app/models/work_packages/scopes/relatable.rb'),
    ('app/models/project.rb', 'https://github.com/opf/openproject/blob/dev/app/models/project.rb'),
    ('app/models/projects/hierarchy.rb', 'https://github.com/opf/openproject/blob/dev/app/models/projects/hierarchy.rb'),
    ('app/models/projects/types.rb', 'https://github.com/opf/openproject/blob/dev/app/models/projects/types.rb'),
    ('app/models/projects/ancestors_from_root.rb', 'https://github.com/opf/openproject/blob/dev/app/models/projects/ancestors_from_root.rb'),
    ('app/models/type.rb', 'https://github.com/opf/openproject/blob/dev/app/models/type.rb'),
    ('app/models/type/attributes.rb', 'https://github.com/opf/openproject/blob/dev/app/models/type/attributes.rb'),
    ('app/models/type/attribute_groups.rb', 'https://github.com/opf/openproject/blob/dev/app/models/type/attribute_groups.rb'),
]
for row_data in ref_rows:
    row = ref_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.save('/home/ubuntu/OpenProject_Reference_Patterns_Full.docx')
print("Vollständiges Word-Dokument erfolgreich erstellt.")
